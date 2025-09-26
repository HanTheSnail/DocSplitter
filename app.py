import streamlit as st
import io
import zipfile
from docx import Document
from docx.shared import Inches
import tempfile
import os

def extract_tables_from_docx(uploaded_file):
    """Extract all tables from a Word document and return them as separate documents"""
    
    # Read the uploaded file
    doc = Document(uploaded_file)
    
    # Check if document has tables
    if not doc.tables:
        return [], "No tables found in the document."
    
    extracted_docs = []
    
    for i, table in enumerate(doc.tables, 1):
        # Create a new document for each table
        new_doc = Document()
        
        # Copy the table to the new document with proper formatting preservation
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        
        # Copy table style if it exists
        if table.style:
            new_table.style = table.style
        
        # Copy each cell's content and formatting exactly
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                new_cell = new_table.cell(row_idx, col_idx)
                
                # Clear the default paragraph in the new cell
                new_cell._element.clear_content()
                
                # Copy all paragraphs from original cell
                for paragraph in cell.paragraphs:
                    # Create new paragraph in the target cell
                    new_para = new_cell.add_paragraph()
                    
                    # Copy paragraph formatting
                    new_para.alignment = paragraph.alignment
                    new_para.style = paragraph.style
                    
                    # Copy all runs with their formatting
                    for run in paragraph.runs:
                        new_run = new_para.add_run(run.text)
                        
                        # Copy all run formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.size = run.font.size
                        new_run.font.name = run.font.name
                        new_run.font.color.rgb = run.font.color.rgb
                        
                        # Copy highlighting/background color if present
                        if run.font.highlight_color:
                            new_run.font.highlight_color = run.font.highlight_color
                
                # Copy cell formatting
                if hasattr(cell, '_element'):
                    # Copy cell borders, shading, etc.
                    new_cell._element.get_or_add_tcPr()
                    if cell._element.tcPr is not None:
                        # This preserves cell-level formatting
                        for child in cell._element.tcPr:
                            new_cell._element.tcPr.append(child)
        
        # Copy table-level formatting
        if hasattr(table._element, 'tblPr') and table._element.tblPr is not None:
            new_table._element.tblPr.clear()
            for child in table._element.tblPr:
                new_table._element.tblPr.append(child)
        
        # Save to memory
        doc_buffer = io.BytesIO()
        new_doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        extracted_docs.append({
            'name': f'table_{i}.docx',
            'content': doc_buffer.getvalue()
        })
    
    return extracted_docs, None

def create_zip_download(docs, original_filename):
    """Create a zip file containing all extracted table documents"""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        base_name = original_filename.rsplit('.', 1)[0]
        
        for doc in docs:
            # Create filename based on original name
            new_filename = f"{base_name}_{doc['name']}"
            zip_file.writestr(new_filename, doc['content'])
    
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def main():
    st.set_page_config(
        page_title="Word Table Splitter",
        page_icon="📊",
        layout="centered"
    )
    
    st.title("📊 Word Document Table Splitter")
    st.markdown("Upload Word documents (.docx) and automatically split each table into separate documents!")
    
    # File uploader
    uploaded_files = st.file_uploader(
        "Choose Word documents",
        type=['docx'],
        accept_multiple_files=True,
        help="Select one or more .docx files containing tables you want to split"
    )
    
    if uploaded_files:
        st.write(f"📁 {len(uploaded_files)} file(s) uploaded")
        
        # Process button
        if st.button("🔄 Split Tables", type="primary"):
            all_results = []
            
            # Progress bar
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for idx, uploaded_file in enumerate(uploaded_files):
                status_text.text(f"Processing: {uploaded_file.name}")
                progress_bar.progress((idx + 1) / len(uploaded_files))
                
                # Extract tables from current file
                extracted_docs, error = extract_tables_from_docx(uploaded_file)
                
                if error:
                    st.error(f"❌ Error with {uploaded_file.name}: {error}")
                    continue
                
                if extracted_docs:
                    all_results.append({
                        'filename': uploaded_file.name,
                        'docs': extracted_docs,
                        'count': len(extracted_docs)
                    })
                    
                    st.success(f"✅ {uploaded_file.name}: Found {len(extracted_docs)} table(s)")
            
            status_text.empty()
            progress_bar.empty()
            
            # Display results and download options
            if all_results:
                st.markdown("---")
                st.subheader("📥 Download Split Documents")
                
                for result in all_results:
                    st.markdown(f"**{result['filename']}** - {result['count']} table(s)")
                    
                    # Create download for individual file results
                    if result['count'] == 1:
                        # Single table - direct download
                        doc = result['docs'][0]
                        st.download_button(
                            label=f"📄 Download {doc['name']}",
                            data=doc['content'],
                            file_name=f"{result['filename'].rsplit('.', 1)[0]}_{doc['name']}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        # Multiple tables - zip download
                        zip_data = create_zip_download(result['docs'], result['filename'])
                        st.download_button(
                            label=f"📦 Download All Tables as ZIP",
                            data=zip_data,
                            file_name=f"{result['filename'].rsplit('.', 1)[0]}_tables.zip",
                            mime="application/zip"
                        )
                
                # Option to download everything as one big zip
                if len(all_results) > 1:
                    st.markdown("---")
                    
                    # Create master zip with all results
                    master_zip_buffer = io.BytesIO()
                    with zipfile.ZipFile(master_zip_buffer, 'w', zipfile.ZIP_DEFLATED) as master_zip:
                        for result in all_results:
                            base_name = result['filename'].rsplit('.', 1)[0]
                            for doc in result['docs']:
                                filename = f"{base_name}_{doc['name']}"
                                master_zip.writestr(filename, doc['content'])
                    
                    master_zip_buffer.seek(0)
                    
                    st.download_button(
                        label="📦 Download All Files as One ZIP",
                        data=master_zip_buffer.getvalue(),
                        file_name="all_split_tables.zip",
                        mime="application/zip",
                        type="primary"
                    )
            
            else:
                st.warning("⚠️ No tables were found in any of the uploaded documents.")
    
    # Instructions
    with st.expander("ℹ️ How to use"):
        st.markdown("""
        1. **Upload** your Word documents (.docx files) using the file uploader above
        2. **Click** the "Split Tables" button to process your documents
        3. **Download** the split table documents:
           - Single table files download directly as .docx
           - Multiple table files download as a ZIP archive
        4. **Extract** ZIP files to access individual table documents
        
        **Note**: Each table from your original document will become a separate Word document,
        preserving the original table structure and basic formatting.
        """)

if __name__ == "__main__":
    main()

if __name__ == "__main__":
    main()
